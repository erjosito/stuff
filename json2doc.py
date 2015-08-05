################################################################################
# Authors: Jose Moreno (based on original script by Camilo Rossi)              #
# josemor@cisco.com                                                            #
#                                                                              #
# August 2015                                                                  #
#                                                                              #
# Takes an input json config of a tenant, that can be obtained by "Saving as"  #
#   from the APIC GUI. It produces a docx document with a text description     #
#                                                                              #
################################################################################

import json
import glob
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import pydot

json_filename = '/Users/josemor/Documents/Python/pod1.json'
docx_filename = '/Users/josemor/Documents/Python/pod1.docx'

document = Document()
obj_styles = document.styles
obj_parstyle = obj_styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
obj_font = obj_parstyle.font
obj_font.size = Pt(11)
obj_font.name = 'Cambria'

unique_keys = set()

# Find the BDs associated to a specific context
def getBdForCtx(fvTenant, ctxName):
    Bds=[]
    for child in fvTenant['children']:
        if 'fvBD' in child:
            Bds.append(child['fvBD']['attributes']['name'])
    return Bds

# Find the L3outs associated to a specific context
def getL3outsForCtx(l3outs, ctxName):
    ctxl3outs=[]
    for l3out in l3outs:
        for child in l3out['children']:
            if 'l3extRsEctx' in child:
                if child['l3extRsEctx']['attributes']['tnFvCtxName'] == ctxName:
                    ctxl3outs.append(l3out['attributes']['name'])
    return ctxl3outs

# Gets the contracts consumed by a specific L3out
def getConsContractsL3out(l3out):
    contractlist = []
    for l3outchild in l3out['children']:
        if 'l3extInstP' in l3outchild:
            for child in l3outchild['l3extInstP']['children']:
                if 'fvRsCons' in child:
                    contractlist.append(child['fvRsCons']['attributes']['tnVzBrCPName'])
    return contractlist

# Gets the subnets announced by a specific L3out
def getExportSubnets(l3out):
    subnetlist = []
    for l3outchild in l3out['children']:
        if 'l3extInstP' in l3outchild:
            for child in l3outchild['l3extInstP']['children']:
                if 'l3extSubnet' in child:
                    if child['l3extSubnet']['attributes']['scope'] == 'export-rtctrl':
                        subnetlist.append(child['l3extSubnet']['attributes']['ip'])
    return subnetlist

# Gets the subnets imported by a specific L3out
def getImportSubnets(l3out):
    subnetlist = []
    for l3outchild in l3out['children']:
        if 'l3extInstP' in l3outchild:
            for child in l3outchild['l3extInstP']['children']:
                if 'l3extSubnet' in child:
                    if child['l3extSubnet']['attributes']['scope'] == 'import-security':
                        subnetlist.append(child['l3extSubnet']['attributes']['ip'])
    return subnetlist

# Get contracts consumed by an EPG
def getConsContractsforEPG (epg):
    contractlist = []
    for child in epg['children']:
        if 'fvRsCons' in child:
            contractlist.append(child['fvRsCons']['attributes']['tnVzBrCPName'])
    return contractlist

# Get contracts provided by an EPG
def getProvContractsforEPG (epg):
    contractlist = []
    for child in epg['children']:
        if 'fvRsProv' in child:
            contractlist.append(child['fvRsProv']['attributes']['tnVzBrCPName'])
    return contractlist

# Gets the BD name in an EPG
def getBdForEPG (epg):
    for child in epg['children']:
        if 'fvRsBd' in child:
            return child['fvRsBd']['attributes']['tnFvBDName']

# Gets the IP addresses (subnets) defined in a BD
def getIpsForBD (bds, bd_name):
    ips=[]
    for bd in bds:
        if bd['attributes']['name'] == bd_name:
            for child in bd['children']:
                if 'fvSubnet' in child:
                    ips.append(child['fvSubnet']['attributes']['ip'])
    return ips


# Load the config file in a dictionary
with open(json_filename) as fd:
    config = json.load(fd)
    fd.close()

# Strip off the initial metadata labels, find out the tenant name
fvTenant = config['imdata'][0]['fvTenant']
document.add_heading('ACI Configuration for tenant ' + fvTenant['attributes']['name'], 0)

# Some lists with JSON code, to make things easier
fvAp = []
fvCtx = []
l3extOut = []
fvBD = []

# Get the list of ANPs, this returns a list so I need to go trough it again.
# To create the document in order I need to divide the objects depending on their class
for child in fvTenant['children']:
    if 'fvAp' in child:
        fvAp.append(child['fvAp'])
    if 'fvCtx' in child:
        fvCtx.append(child['fvCtx'])
    if 'l3extOut' in child:
        l3extOut.append(child['l3extOut'])
    if 'fvBD' in child:
        fvBD.append(child['fvBD'])

# Networking info (VRFs, BDs, L3outs)
document.add_heading('Networking configuration', level=1)

# Print a brief description of the private networks
document.add_paragraph('', style='CodeStyle')
document.add_heading('Private networks (VRFs)', level=2)
document.add_paragraph('', style='CodeStyle')
if fvCtx.__len__() == 0:
    document.add_paragraph('This tenant has no private networks (VRFs) defined, it is probably using a private network defined in the common tenant.', style='CodeStyle')
else:
    document.add_paragraph('This tenant has the following private networks (VRFs) defined:', style='CodeStyle')
    for vrf in fvCtx:
        paragraph=vrf['attributes']['name']
        document.add_paragraph(paragraph, style='CodeStyle')
        # Bridge domains
        bds = getBdForCtx(fvTenant,vrf['attributes']['name'])
        if bds.__len__() == 0:
            paragraph = 'No bridge domains associated to this VRF.'
            document.add_paragraph(paragraph, style='CodeStyle')
        else:
            paragraph = 'The following bridge domains are defined in this VRF: '
            document.add_paragraph(paragraph, style='CodeStyle')
            for bd in bds:
                paragraph = bd
                ips = getIpsForBD(fvBD, bd)
                if ips.__len__() == 0:
                    paragraph = paragraph + '. No subnets are configured in this bridge domain'
                else:
                    paragraph = paragraph + '. The following subnets are configured in this bridge domain: '
                    for ip in ips:
                        paragraph = paragraph + ip + ", "
                    # Remove last comma
                    paragraph = paragraph[0:paragraph.__len__()-2]
                document.add_paragraph(paragraph, style='ListBullet')
        # L3 outs
        l3outs = getL3outsForCtx(l3extOut,vrf['attributes']['name'])
        if l3outs.__len__() == 0:
            paragraph = 'No external IP connections associated to this VRF'
            document.add_paragraph(paragraph, style='CodeStyle')
        else:
            paragraph = 'The following external IP connections are associated to his VRF: '
            document.add_paragraph(paragraph, style='CodeStyle')
            for l3out in l3outs:
                document.add_paragraph(l3out, style='ListBullet')
        # Create image
        graph = pydot.Dot(graph_type='graph')
        for bd in bds:
            edge = pydot.Edge('VRF '+ vrf['attributes']['name'], 'BD ' + bd)
            graph.add_edge(edge)
            ips = getIpsForBD(fvBD, bd)
            for ip in ips:
                edge = pydot.Edge('BD ' + bd, ip)
                graph.add_edge(edge)
        for l3out in l3outs:
            edge = pydot.Edge('VRF ' + vrf['attributes']['name'], 'L3out ' + l3out)
            graph.add_edge(edge)
        graph.write_png(vrf['attributes']['name'] + '.png')
        document.add_picture(vrf['attributes']['name'] + '.png')


# Print a brief description of the L3outs
document.add_paragraph('', style='CodeStyle')
document.add_heading('External IP connections', level=2)
document.add_paragraph(' ', style='CodeStyle')
if l3extOut.__len__() == 0:
    document.add_paragraph('This tenant has no external IP connections (L3) defined: it is either isolated, or using a L2 connection to the outside world.', style='CodeStyle')
else:
    document.add_paragraph('This tenant has the following external IP connections defined:', style='CodeStyle')
    for l3out in l3extOut:
        paragraph='-   ' + l3out['attributes']['name']
        importsubnets = getImportSubnets(l3out)
        exportsubnets = getExportSubnets(l3out)
        contracts=getConsContractsL3out(l3out)
        if importsubnets.__len__() == 0:
            paragraph = paragraph + '. No subnets are configured to be received by this L3 connection'
        else:
            paragraph = paragraph + '. The following subnets are configured to be received by this L3 connection: '
            for subnet in importsubnets:
                paragraph = paragraph + subnet + ", "
            # Remove last comma
            paragraph = paragraph[0:paragraph.__len__()-2]
        if exportsubnets.__len__() == 0:
            paragraph = paragraph + '. No subnets are configured to be announced by this IP connection'
        else:
            paragraph = paragraph + '. The following subnets are configured to be announced by this IP connection: '
            for subnet in exportsubnets:
                paragraph = paragraph + subnet + ", "
            # Remove last comma
            paragraph = paragraph[0:paragraph.__len__()-2]
        if contracts.__len__() == 0:
            paragraph = paragraph + '. No contracts are consumed by this IP connection'
        else:
            paragraph = paragraph + '. The following contracts are consumed by this IP connection: '
            for contract in contracts:
                paragraph = paragraph + contract + ", "
            # Remove last comma
            paragraph = paragraph[0:paragraph.__len__()-2]
        paragraph = paragraph + '.'
        document.add_paragraph(paragraph, style='CodeStyle')

# Print one section for each ANP
document.add_paragraph(' ', style='CodeStyle')
document.add_heading('Application Network Profiles (ANPs)', level=1)
document.add_paragraph(' ', style='CodeStyle')
if fvCtx.__len__() == 0:
    document.add_paragraph('This tenant has no application network profiles defined.', style='CodeStyle')
else:
    document.add_paragraph('The following application network profiles are defined in this tenant:', style='CodeStyle')
    for anp in fvAp:
        document.add_paragraph('', style='CodeStyle')
        document.add_heading('ANP ' + anp['attributes']['name'], level=2)
        fvAEPg = []
        for child in anp['children']:
            if 'fvAEPg' in child:
                fvAEPg.append(child['fvAEPg'])
        if fvAEPg.__len__() == 0:
            document.add_paragraph('This application network profile has no End Point Groups (EPGs) defined.', style='CodeStyle')
        else:
            document.add_paragraph('This application network profile has the following End Point Groups (EPGs) defined:', style='CodeStyle')
            for epg in fvAEPg:
                paragraph='-   ' + epg['attributes']['name'] + ':'
                # BD
                paragraph = paragraph + ' This EPG is associated to Bridge Domain ' + getBdForEPG(epg) + '.'
                # Contracts
                cons = getConsContractsforEPG(epg)
                prov = getProvContractsforEPG(epg)
                if cons.__len__() == 0:
                    paragraph = paragraph + ' No contracts are consumed by this EPG'
                else:
                    paragraph = paragraph + ' The following contracts are consumed by this EPG: '
                    for contract in cons:
                        paragraph = paragraph + contract + ", "
                    # Remove last comma
                    paragraph = paragraph[0:paragraph.__len__()-2]
                if prov.__len__() == 0:
                    paragraph = paragraph + '. No contracts are provided by this EPG'
                else:
                    paragraph = paragraph + '. The following contracts are provided by this EPG: '
                    for contract in prov:
                        paragraph = paragraph + contract + ", "
                    # Remove last comma
                    paragraph = paragraph[0:paragraph.__len__()-2]
                document.add_paragraph(paragraph, style='CodeStyle')
        # Create image
        graph = pydot.Dot(graph_type='digraph')
        epg_nodes = {}
        epgs=[]
        # Build a matrix with EPG names and a dictionary with EPG names and nodes
        for epg in fvAEPg:
            epg_name=epg['attributes']['name']
            epg_nodes[epg_name] = pydot.Node("EPG " + epg_name, style="filled", fillcolor="white")
            epgs.append(epg_name)
        for epg in epgs:
            graph.add_node(epg_nodes[epg])
        # For each EPG...
        for srcepg in fvAEPg:
            srcepg_name = srcepg['attributes']['name']
            # ... and each provided contract in that EPG
            prov = getProvContractsforEPG(srcepg)
            for srccontract in prov:
                # Try to find another EPG...
                for dstepg in fvAEPg:
                    dstepg_name = dstepg['attributes']['name']
                    cons = getConsContractsforEPG(dstepg)
                    # ...that consumes that contract
                    for dstcontract in cons:
                        if srccontract == dstcontract and srcepg_name != dstepg_name:
                            graph.add_edge(pydot.Edge(epg_nodes[srcepg_name], epg_nodes[dstepg_name], label=srccontract))
                # Or try to find a L3out
                for l3out in l3extOut:
                    l3out_name = l3out['attributes']['name']
                    cons = getConsContractsL3out(l3out)
                    # ...that consumes that contract
                    for dstcontract in cons:
                        if srccontract == dstcontract:
                            l3outnode = pydot.Node("L3out " + l3out_name, style="filled", fillcolor="white")
                            graph.add_edge(pydot.Edge(epg_nodes[srcepg_name], l3outnode, label=srccontract))
        graph.write_png(anp['attributes']['name'] + '.png')
        document.add_picture(anp['attributes']['name'] + '.png')



print "Saving document..."
document.save(docx_filename)
